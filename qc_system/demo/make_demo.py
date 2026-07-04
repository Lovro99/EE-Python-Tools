"""Generator demo projekta s NAMJERNIM neskladima.

Kreira tri datoteke kakve nastaju u stvarnom radu:
  - bilanca_snaga.xlsx   (Excel bilanca — "izvor istine")
  - export_dwg.csv       (kao export atributa blokova iz AutoCAD-a)
  - tehnicki_opis.docx   (Word opis s tablicom krugova)

Ugradene greske koje QC mora naci:
  1. RO-1: snaga u DWG-u 12 kW, u Excelu i Wordu 15 kW    -> GRESKA
  2. RO-3: kabel u Wordu 5x4, drugdje 5x6                 -> GRESKA
  3. S2.1: zastita postoji samo u Excelu                  -> UPOZORENJE
  4. RO-2: u Excelu "10,0 kW", u DWG "10" — ISTO, ne smije
     biti prijavljeno (test normalizacije)
"""

import csv
import shutil
from pathlib import Path

from docx import Document
from openpyxl import Workbook

# (oznaka, snaga, struja, kabel, zastita)
KRUGOVI = [
    ("RO-1", "15,0", "23", "NYY-J 5×6", "C25"),
    ("RO-2", "10,0", "16", "NYY-J 5×2,5", "C16"),
    ("RO-3", "7,5", "12", "NYY-J 5×6", "C13"),
    ("S2.1", "3,0", "5", "NYM-J 3×1,5", "B10"),
]


def _make_excel(path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Bilanca"
    ws["A1"] = "BILANCA SNAGA — DEMO PROJEKT"
    headers = ["Oznaka", "Opis", "P (kW)", "In (A)", "Kabel", "Zaštita"]
    for col, h in enumerate(headers, start=1):
        ws.cell(row=2, column=col, value=h)
    for i, (tag, kw, a, kabel, zastita) in enumerate(KRUGOVI, start=3):
        ws.cell(row=i, column=1, value=tag)
        ws.cell(row=i, column=2, value="Trošilo " + tag)
        ws.cell(row=i, column=3, value=kw + " kW")
        ws.cell(row=i, column=4, value=a)
        ws.cell(row=i, column=5, value=kabel)
        ws.cell(row=i, column=6, value=zastita)
    wb.save(path)


def _make_dwg_csv(path):
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow(["CIRCUIT_LABEL", "SNAGA", "STRUJA", "KABEL", "ZASTITA"])
        for tag, kw, a, kabel, zastita in KRUGOVI:
            if tag == "RO-1":
                kw = "12"          # GRESKA 1: nacrt nije azuriran (12 vs 15)
            elif tag == "RO-2":
                kw = "10"          # test normalizacije: "10" == "10,0 kW"
            if tag == "S2.1":
                zastita = ""       # UPOZORENJE: zastita samo u Excelu
            # nacrt pise kabel bez × i s tockom — test normalizacije
            kabel = kabel.replace("×", "x").replace(",", ".")
            w.writerow([tag, kw, a, kabel, zastita])


def _make_word(path):
    doc = Document()
    doc.add_heading("Tehnički opis — DEMO PROJEKT", level=1)
    doc.add_paragraph(
        "Napajanje trošila izvodi se prema tablici strujnih krugova:"
    )
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Oznaka"
    hdr[1].text = "Snaga (kW)"
    hdr[2].text = "Kabel"
    hdr[3].text = "Zaštita"
    for tag, kw, a, kabel, zastita in KRUGOVI:
        if tag == "RO-3":
            kabel = "NYY-J 5×4"    # GRESKA 2: opis ima stari presjek
        if tag == "S2.1":
            zastita = ""           # zastita nedostaje i u opisu
        row = table.add_row().cells
        row[0].text = tag.lower()  # test normalizacije oznake (ro-1 == RO-1)
        row[1].text = kw
        row[2].text = kabel
        row[3].text = zastita
    doc.save(path)


def make_demo_project(demo_dir):
    demo_dir = Path(demo_dir)
    if demo_dir.exists():
        shutil.rmtree(demo_dir)
    demo_dir.mkdir(parents=True)
    _make_excel(demo_dir / "bilanca_snaga.xlsx")
    _make_dwg_csv(demo_dir / "export_dwg.csv")
    _make_word(demo_dir / "tehnicki_opis.docx")
    return demo_dir


if __name__ == "__main__":
    make_demo_project(Path(__file__).parent.parent / "demo_projekt")
