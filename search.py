import os
import glob
from pathlib import Path
import pickle
from collections import defaultdict
import tkinter as tk
from tkinter import filedialog, messagebox
import time
import re
import io
import zipfile
from docx import Document

class DocxIndexer:
    def __init__(self):
        self.word_index = defaultdict(list)
        self.full_content = {}
    
    def build_index(self, root):
        patterns = ["**/*.docx", "**/*.docm"]
        total_files = 0
        
        print("Počinjem indeksiranje...")
        start_time = time.time()
        
        for pattern in patterns:
            search_path = str(Path(root) / pattern)
            for path in glob.glob(search_path, recursive=True):
                if Path(path).name.startswith("~"):
                    continue
                
                total_files += 1
                print(f"Obrađujem ({total_files}): {os.path.basename(path)}")
                
                page_data = self.extract_content_with_pages(path)
                if page_data:
                    self.full_content[path] = page_data
                    for page_info in page_data:
                        words = self.tokenize(page_info['text'])
                        for word in words:
                            if not any(m['path'] == path and page_info['page'] in m['pages'] 
                                     for m in self.word_index[word]):
                                self.word_index[word].append({
                                    'path': path, 
                                    'pages': [page_info['page']]
                                })
        
        elapsed = time.time() - start_time
        print(f"✅ Indeksiranje završeno: {total_files} datoteka ({elapsed:.1f}s)")
    
    def extract_content_with_pages(self, path):
        """
        Popravljeno za .docm: Otvara datoteku kao ZIP, mijenja content-type 
        u memoriji i servira je docx biblioteci kao običan .docx
        """
        try:
            # Ako je .docm, moramo ga 'presvući' u .docx u memoriji
            if path.lower().endswith('.docm'):
                memory_file = io.BytesIO()
                with zipfile.ZipFile(path, 'r') as zin:
                    with zipfile.ZipFile(memory_file, 'w') as zout:
                        for item in zin.infolist():
                            buffer = zin.read(item.filename)
                            # Zamijeni referencu na macroEnabled u običan document
                            if item.filename == '[Content_Types].xml':
                                buffer = buffer.replace(
                                    b'application/vnd.ms-word.document.macroEnabled.main+xml',
                                    b'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'
                                )
                            zout.writestr(item, buffer)
                memory_file.seek(0)
                doc = Document(memory_file)
            else:
                doc = Document(path)
                
            return self._process_doc_obj(doc)
            
        except Exception as e:
            print(f"❌ Neuspjelo otvaranje {os.path.basename(path)}: {e}")
            return None

    def _process_doc_obj(self, doc):
        page_data = []
        total_chars = 0
        total_paragraphs = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                chars = len(text)
                total_chars += chars
                total_paragraphs += 1
                # Aproksimacija stranice
                page_num = max(1, total_chars // 2100 + total_paragraphs // 28)
                
                page_data.append({
                    'page': page_num,
                    'text': text.lower()
                })
                
                if total_chars > 3000:
                    total_chars = 0
                    total_paragraphs = 0
        return page_data

    def tokenize(self, text):
        words = set()
        # Zadržava hrvatska slova, brojeve i specifične znakove poput 3x, 400V
        for word in text.split():
            word_clean = re.sub(r'[.,;:!?()\[\]{}„"\'–—]', '', word.lower())
            if len(word_clean) > 1:
                words.add(word_clean)
        return words

    def save_index(self, filepath):
        data = {'full_content': self.full_content, 'word_index': dict(self.word_index)}
        with open(filepath, 'wb') as f:
            pickle.dump(data, f)
        print(f"💾 Indeks spremljen: {filepath}")
    
    def load_index(self, filepath):
        with open(filepath, 'rb') as f:
            data = pickle.load(f)
        self.full_content = data.get('full_content', {})
        self.word_index = defaultdict(list, data.get('word_index', {}))
        print(f"📂 Indeks učitan: {len(self.full_content)} datoteka")

    def search_flexible(self, query):
        query = query.lower().strip()
        if ' ' not in query:
            matches = self.word_index.get(query, [])
            if matches: return matches
            pattern = f'.*{re.escape(query)}.*'
            all_matches = []
            for word, entries in self.word_index.items():
                if re.match(pattern, word): all_matches.extend(entries)
            return all_matches
        return self.search_full_content(query)
    
    def search_full_content(self, query):
        matches = []
        query_lower = query.lower()
        for path, page_data in self.full_content.items():
            for page_info in page_data:
                if query_lower in page_info['text']:
                    if not any(m['path'] == path and page_info['page'] in m['pages'] for m in matches):
                        matches.append({'path': path, 'pages': [page_info['page']]})
        return matches

    def get_context_around_word(self, path, page_num, search_term):
        if path not in self.full_content: return []
        contexts = []
        search_words = search_term.lower().split()
        for page_info in self.full_content[path]:
            if page_info['page'] == page_num:
                words = page_info['text'].split()
                for i, word in enumerate(words):
                    word_clean = re.sub(r'[.,;:!?()\[\]{}„"\'\-–—]', '', word.lower())
                    if all(term in word_clean for term in search_words):
                        start = max(0, i-3)
                        end = min(len(words), i+4)
                        contexts.append(' '.join(words[start:end]))
        return contexts

class SearchApp:
    def __init__(self, indexer):
        self.indexer = indexer
        self.root = tk.Tk()
        self.root.title("🔍 DOCX/DOCM Pretraživač v3.4")
        self.root.geometry("1100x800")
        self.setup_ui()
    
    def setup_ui(self):
        # Header
        header = tk.Frame(self.root, bg="#2d3748", height=80)
        header.pack(fill="x")
        tk.Label(header, text="🔍 DOCX & DOCM Pretraživač", font=("Arial", 16, "bold"), bg="#2d3748", fg="white").pack(pady=10)
        
        # Search area
        s_frame = tk.Frame(self.root)
        s_frame.pack(fill="x", padx=20, pady=15)
        
        self.search_entry = tk.Entry(s_frame, font=("Arial", 14), relief="solid", bd=2)
        self.search_entry.pack(fill="x", pady=5)
        self.search_entry.bind("<Return>", lambda e: self.do_search())
        
        btn_box = tk.Frame(s_frame)
        btn_box.pack()
        tk.Button(btn_box, text="🔍 TRAŽI", command=self.do_search, bg="#48bb78", fg="white", font=("Arial", 11, "bold"), padx=20).pack(side="left", padx=5)
        tk.Button(btn_box, text="🗑️ OBRIŠI", command=self.clear_results, bg="#ed8936", fg="white", font=("Arial", 11, "bold"), padx=20).pack(side="left", padx=5)
        
        # Results area
        r_frame = tk.Frame(self.root)
        r_frame.pack(fill="both", expand=True, padx=20, pady=10)
        self.results_text = tk.Text(r_frame, wrap=tk.WORD, font=("Consolas", 11), bg="#f8f9fa")
        sb = tk.Scrollbar(r_frame, command=self.results_text.yview)
        self.results_text.configure(yscrollcommand=sb.set)
        self.results_text.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        
        # Footer
        footer = tk.Frame(self.root)
        footer.pack(fill="x", padx=20, pady=10)
        tk.Button(footer, text="📋 KOPIRAJ REZULTATE", command=self.copy_results, bg="#4299e1", fg="white", font=("Arial", 10, "bold"), padx=15).pack(side="left")
        tk.Button(footer, text="❌ IZLAZ", command=self.root.quit, bg="#f56565", fg="white", font=("Arial", 10, "bold"), padx=15).pack(side="right")

    def do_search(self):
        query = self.search_entry.get().strip()
        if not query: return
        self.results_text.delete("1.0", tk.END)
        matches = self.indexer.search_flexible(query)
        self.show_results(matches, query)

    def show_results(self, matches, query):
        if not matches:
            self.results_text.insert("1.0", f"❌ Nema rezultata za: '{query}'")
            return
        
        file_pages = defaultdict(list)
        total_hits = 0
        for m in matches:
            if isinstance(m, dict):
                path = m['path']
                file_pages[path].extend(m['pages'])
                total_hits += len(m['pages'])

        res_head = f"🔍 UPIT: '{query}' | 📊 Pojavljivanja: {total_hits} | 📄 Datoteka: {len(file_pages)}\n"
        res_head += "="*70 + "\n\n"
        self.results_text.insert(tk.END, res_head)

        for path, pages in sorted(file_pages.items()):
            fname = os.path.basename(path)
            self.results_text.insert(tk.END, f"📄 {fname}\n")
            self.results_text.insert(tk.END, f"   📍 {path}\n")
            for p in sorted(set(pages)):
                contexts = self.indexer.get_context_around_word(path, p, query)
                if contexts:
                    ctx_str = " | ".join(f"«{c}»" for c in contexts[:3])
                    self.results_text.insert(tk.END, f"   📖 Str. {p}: {ctx_str}\n")
                else:
                    self.results_text.insert(tk.END, f"   📖 Str. {p}\n")
            self.results_text.insert(tk.END, "\n")

    def clear_results(self):
        self.results_text.delete("1.0", tk.END)
        self.search_entry.delete(0, tk.END)

    def copy_results(self):
        txt = self.results_text.get("1.0", tk.END).strip()
        if txt:
            self.root.clipboard_clear()
            self.root.clipboard_append(txt)
            messagebox.showinfo("Kopirano", "Rezultati su kopirani u clipboard!")

    def run(self):
        """Ova metoda je nedostajala i uzrokovala je AttributeError"""
        self.root.mainloop()

def main():
    indexer = DocxIndexer()
    
    # Prozor za odabir
    startup_root = tk.Tk()
    startup_root.withdraw()
    
    ans = messagebox.askyesno("Indeksiranje", "Želite li učitati postojeću bazu (.pkl)?\n(Kliknite 'No' za novo indeksiranje mape)")
    
    if ans:
        idx_file = filedialog.askopenfilename(filetypes=[("Pickle datoteka", "*.pkl")])
        if idx_file:
            indexer.load_index(idx_file)
        else:
            startup_root.destroy()
            return
    else:
        folder = filedialog.askdirectory(title="Odaberi mapu za indeksiranje (.docx i .docm)")
        if not folder:
            startup_root.destroy()
            return
        
        indexer.build_index(folder)
        save_p = filedialog.asksaveasfilename(
            title="Spremi bazu za buduće pretrage",
            defaultextension=".pkl", 
            filetypes=[("Pickle datoteka", "*.pkl")]
        )
        if save_p:
            indexer.save_index(save_p)
    
    startup_root.destroy()
    app = SearchApp(indexer)
    app.run()

if __name__ == "__main__":
    main()