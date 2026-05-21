from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# --- Styles ---
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Calibri'
h1_style.font.size = Pt(20)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Calibri'
h2_style.font.size = Pt(14)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

h3_style = doc.styles['Heading 3']
h3_style.font.name = 'Calibri'
h3_style.font.size = Pt(12)
h3_style.font.bold = True
h3_style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_info_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run_label = p.add_run(label + ': ')
    run_label.bold = True
    run_label.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run_label.font.size = Pt(10.5)
    run_value = p.add_run(value)
    run_value.font.size = Pt(10.5)


def add_command_box(doc, commands):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF3FB')
    cell._tc.tcPr.append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(commands)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph()


# ==================== NASLOVNICA ====================
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('AutoLisp Programi')
title_run.font.name = 'Calibri'
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('Dokumentacija i upute za korištenje')
sub_run.font.name = 'Calibri'
sub_run.font.size = Pt(16)
sub_run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
add_horizontal_line(doc)
doc.add_paragraph()

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc_p.add_run('Zbirka prilagođenih AutoLisp programa za AutoCAD\nza automatizaciju crteža, upravljanje blokovima,\ntekstom, rasporedom i integraciju s Excelom.')
desc_run.font.name = 'Calibri'
desc_run.font.size = Pt(12)
desc_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run('Verzija: svibanj 2026.')
date_run.font.name = 'Calibri'
date_run.font.size = Pt(11)
date_run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

# Prijelom stranice
doc.add_page_break()

# ==================== SADRŽAJ ====================
toc_heading = doc.add_heading('Sadržaj', level=1)
toc_heading.paragraph_format.space_after = Pt(12)

def add_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

add_toc(doc)
doc.add_page_break()

# ==================== UVOD ====================
doc.add_heading('Uvod', level=1)
intro = doc.add_paragraph(
    'Ovaj dokument sadrži dokumentaciju i upute za korištenje svih prilagođenih AutoLisp (.lsp) '
    'programa koji se koriste u AutoCAD-u. Programi su razvijeni za automatizaciju svakodnevnih '
    'zadataka poput upravljanja blokovima i atributima, obrade teksta, upravljanja layoutima, '
    'integracije s programom Microsoft Excel te specifičnih zadataka vezanih uz električne nacrte '
    'i solarne sustave.'
)
intro.paragraph_format.space_after = Pt(8)

intro2 = doc.add_paragraph(
    'Programi su poredani abecednim redom. Za svaki program navedeni su: opis funkcionalnosti, '
    'naredbe kojima se program aktivira te kratke upute za korištenje.'
)
intro2.paragraph_format.space_after = Pt(8)

note = doc.add_paragraph()
note_run_bold = note.add_run('Napomena: ')
note_run_bold.bold = True
note_run_bold.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
note_run = note.add_run(
    'Svi programi učitavaju se automatski pri pokretanju AutoCAD-a putem datoteke ACADDOC.lsp. '
    'Naredbe se unose u naredbenu liniju AutoCAD-a (Command Line).'
)

doc.add_page_break()

# ==================== KATEGORIJE ====================
doc.add_heading('Kategorije programa', level=1)

categories = [
    ('Upravljanje blokovima i atributima',
     'BatchAttributeEditorV1-5, BurstUpgradedV1-7, ChangeBlockBasePointV1-5, CountAttributeValues, '
     'DBCountV1-1, ExtractNestedBlockV1-2, InsertElements, StealV1-8, insertBlock, simpleCount'),
    ('Upravljanje tekstom',
     'BoxTextV1-2, CopySwapTextV1-8, GrTextV1-1, MTVATT, ModifyTextValues, NumIncV4-1, '
     'TCountV1-1, Text2MTextV2-0, tlen'),
    ('Upravljanje layoutima',
     'RenumberLayoutsV1-2, TabSortV2-2, VportsToggle, createlayoutV2, setLayoutOrderV5, SetLayoutTitles'),
    ('Električni nacrti i VDS',
     'CBnumbering, CountCB, VDS, VDSnum2, setCBelements'),
    ('Solarni sustavi',
     'MarenConnectPro, panelRenum, sumPowerV2'),
    ('Integracija s Excelom',
     'ExportLayoutsToExcel, GetExcel, SetFieldsValue'),
    ('Pomoćni alati i ostalo',
     'ACADDOC, ACADLT, Autoloader, CleanUp, LMpopup, Mastermind, makemore'),
]

for cat_name, cat_programs in categories:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run_bold = p.add_run(cat_name + ': ')
    run_bold.bold = True
    run_bold.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run_text = p.add_run(cat_programs)

doc.add_page_break()

# ==================== PROGRAMI ====================
doc.add_heading('Programi', level=1)

programs = [
    {
        'name': 'ACADDOC.lsp',
        'category': 'Konfiguracija',
        'command': '— (automatsko učitavanje)',
        'author': 'Prilagođeno',
        'description': (
            'Ova datoteka služi kao glavni autoloader koji se izvršava automatski pri svakom pokretanju AutoCAD-a. '
            'Njena uloga je da učita (autoload) sve prilagođene LSP programe i registrira njihove naredbe kako bi '
            'bile dostupne u AutoCAD sesiji.'
        ),
        'usage': [
            'Datoteka se ne pokreće ručno – AutoCAD je automatski izvršava pri svakom pokretanju.',
            'Da biste dodali novi program u autoloader, dodajte odgovarajući (autoload ...) unos u ovu datoteku.',
            'Putanje do LSP datoteka trebaju biti ispravno postavljene u AutoCAD support path.',
        ],
        'notes': 'Ako neka naredba nije dostupna, provjerite je li odgovarajući unos dodan u ACADDOC.lsp.',
    },
    {
        'name': 'ACADLT.lsp',
        'category': 'Konfiguracija',
        'command': '— (automatsko učitavanje)',
        'author': 'Prilagođeno',
        'description': (
            'Alternativna verzija autoloader datoteke namijenjena za drugačiju konfiguraciju AutoCAD-a. '
            'Funkcionira jednako kao ACADDOC.lsp, ali se koristi u posebnim okolnostima ili s drugačijim postavkama.'
        ),
        'usage': [
            'Datoteka se ne pokreće ručno – AutoCAD je koristi automatski.',
            'Primijenite ju kao zamjenu za ACADDOC.lsp kada je potrebna alternativna konfiguracija.',
        ],
        'notes': None,
    },
    {
        'name': 'Autoloader.lsp',
        'category': 'Pomoćni alati',
        'command': 'AUTOLOADER',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Program koji automatski generira autoload izraze za sve LSP datoteke u odabranom direktoriju. '
            'Rezultat se zapisuje u tekstualnu datoteku koja se može kopirati u ACADDOC.lsp kako bi se '
            'automatski učitali svi programi iz odabranog direktorija.'
        ),
        'usage': [
            'Unesite naredbu AUTOLOADER u naredbenu liniju.',
            'Odaberite direktorij koji sadrži LSP datoteke.',
            'Program će generirati tekstualnu datoteku s autoload izrazima.',
            'Sadržaj generirane datoteke kopirajte u ACADDOC.lsp.',
        ],
        'notes': None,
    },
    {
        'name': 'BatchAttributeEditorV1-5.lsp',
        'category': 'Blokovi i atributi',
        'command': 'BAE',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Napredni uređivač atributa koji omogućuje skupno uređivanje vrijednosti atributa u više blokova '
            'istovremeno. Podržava pretraživanje i zamjenu, filtriranje po imenu bloka i tagu atributa, '
            'te uvoz i izvoz podataka.'
        ),
        'usage': [
            'Unesite naredbu BAE u naredbenu liniju.',
            'Otvorit će se dijaloški prozor s listom blokova i njihovih atributa.',
            'Odaberite atribute koje želite urediti.',
            'Unesite novu vrijednost ili koristite funkciju "Find & Replace".',
            'Potvrdite promjene klikom na "Apply" ili "OK".',
        ],
        'notes': 'Program podržava rad s više dokumenata i složene strukture ugniježđenih blokova.',
    },
    {
        'name': 'BoxTextV1-2.lsp',
        'category': 'Tekst',
        'command': 'BT',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Program koji crta pravokutni okvir (2D polilinija) oko odabranog teksta ili MText objekta. '
            'Korisnik može definirati odmak (offset) između teksta i okvira. Korisno za isticanje '
            'važnih tekstualnih oznaka u nacrtu.'
        ),
        'usage': [
            'Unesite naredbu BT u naredbenu liniju.',
            'Odaberite jedan ili više tekst/MText objekata.',
            'Unesite vrijednost odmaka (offset) ili pritisnite Enter za zadanu vrijednost.',
            'Program će nacrtati pravokutne okvire oko svakog odabranog teksta.',
        ],
        'notes': 'Okviri se crtaju na aktivnom layeru s aktivnim postavkama crte.',
    },
    {
        'name': 'BurstUpgradedV1-7.lsp',
        'category': 'Blokovi i atributi',
        'command': 'BURST',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Nadograđena verzija AutoCAD naredbe BURST. Eksplodira blokove i pretvara vidljive atribute '
            'u Text ili MText objekte, zadržavajući pri tome sve vizualne svojstva (font, veličina, boja, sloj). '
            'Podržava i ugniježđene (nested) blokove te dinamičke blokove.'
        ),
        'usage': [
            'Unesite naredbu BURST u naredbenu liniju.',
            'Odaberite blokove koje želite eksplodirati.',
            'Pritisnite Enter za potvrdu selekcije.',
            'Program će eksplodirati blokove i pretvoriti atribute u tekst objekte.',
        ],
        'notes': 'Za razliku od standardne AutoCAD BURST naredbe, ova verzija čuva sva vizualna svojstva atributa.',
    },
    {
        'name': 'CBnumbering.lsp',
        'category': 'Električni nacrti',
        'command': 'CBN',
        'author': 'Prilagođeno',
        'description': (
            'Program za numeriranje elemenata sklopnih blokova (CB – Circuit Breaker). Podržava prefikse '
            'F, F0, Q i N s mogućnošću podešavanja početnog broja numeriranja. '
            'Koristi se za automatsko označavanje zaštitnih komponenti u elektroenergetskim shemama.'
        ),
        'usage': [
            'Unesite naredbu CBN u naredbenu liniju.',
            'Odaberite prefiks (F, F0, Q ili N).',
            'Unesite početni broj numeriranja.',
            'Odabirom blokova u nacrtu program će automatski dodijeliti oznake.',
        ],
        'notes': None,
    },
    {
        'name': 'ChangeBlockBasePointV1-5.lsp',
        'category': 'Blokovi i atributi',
        'command': 'CBP',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Program koji mijenja baznu točku definicije bloka. Korisnik može odabrati hoće li vizualni položaj '
            'bloka u nacrtu ostati nepromijenjen ili hoće li se točka umetanja zadržati. '
            'Korisno za ispravljanje pogrešno postavljenih baznih točaka blokova.'
        ),
        'usage': [
            'Unesite naredbu CBP u naredbenu liniju.',
            'Odaberite blok čiju baznu točku želite promijeniti.',
            'Odredite novu baznu točku klikom u nacrtu ili unosom koordinata.',
            'Odaberite opciju zadržavanja vizualnog položaja.',
        ],
        'notes': 'Promjena bazne točke utječe na sve instance bloka u nacrtu.',
    },
    {
        'name': 'CleanUp.lsp',
        'category': 'Pomoćni alati',
        'command': 'CUP',
        'author': 'Prilagođeno',
        'description': (
            'Uslužni program za čišćenje i organizaciju nacrta. Omogućuje odabir blokova, šrafura ili '
            'postavljanje boje 250 na odabrane objekte. Koristi se za brzo ujednačavanje izgleda elemenata u nacrtu.'
        ),
        'usage': [
            'Unesite naredbu CUP u naredbenu liniju.',
            'Odaberite željenu operaciju iz izbornika.',
            'Slijedite upute na zaslonu.',
        ],
        'notes': None,
    },
    {
        'name': 'CopySwapTextV1-8.lsp',
        'category': 'Tekst',
        'command': 'CST / SWAPTEXT',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Program koji omogućuje kopiranje ili zamjenu sadržaja teksta između više tekst objekata '
            '(Text, MText, atributi, dimenzije, multivođice). Podržava prijenos formatiranja i rad '
            's više odabranih objekata istovremeno.'
        ),
        'usage': [
            'Unesite naredbu CST (kopiranje) ili SWAPTEXT (zamjena).',
            'Odaberite izvorni tekst objekt.',
            'Odaberite odredišni tekst objekt ili objekte.',
            'Program će kopirati ili zamijeniti sadržaj teksta.',
        ],
        'notes': 'Opcija formatiranja kontrolira prenosi li se i stil teksta zajedno s tekstualnim sadržajem.',
    },
    {
        'name': 'CountAttributeValues.lsp',
        'category': 'Blokovi i atributi',
        'command': 'CAV',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Program koji prebroji pojave vrijednosti atributa u atributiranim blokovima i prikaže '
            'rezultate u obliku AutoCAD tablice. Omogućuje filtriranje po imenu bloka i tagu atributa.'
        ),
        'usage': [
            'Unesite naredbu CAV u naredbenu liniju.',
            'Odaberite blokove koje želite analizirati ili koristite "Select All".',
            'Navedite ime bloka i tag atributa koji se broji.',
            'Odredite mjesto umetanja rezultatne tablice u nacrtu.',
        ],
        'notes': None,
    },
    {
        'name': 'CountCB.lsp',
        'category': 'Električni nacrti',
        'command': 'CCB',
        'author': 'Prilagođeno',
        'description': (
            'Program koji analizira komponente sklopnih blokova u nacrtu i generira CSV izvještaj '
            'sa specifikacijama. Automatski prikuplja podatke o vrstama i brojevima komponenti '
            'za potrebe dokumentacije i narudžbi materijala.'
        ),
        'usage': [
            'Unesite naredbu CCB u naredbenu liniju.',
            'Odaberite sklopne blokove za analizu.',
            'Navedite putanju i naziv CSV datoteke za izvoz.',
            'Program će generirati CSV s popisom komponenti i njihovim specifikacijama.',
        ],
        'notes': None,
    },
    {
        'name': 'createlayoutV2.lsp',
        'category': 'Layouti',
        'command': 'CRL',
        'author': 'Prilagođeno',
        'description': (
            'Program koji automatski kreira layoute na osnovi odabranih predložaka blokova. '
            'Automatski postavlja page setup i upravlja atributima titlebloka. '
            'Ubrzava pripremu nacrta s većim brojem listova.'
        ),
        'usage': [
            'Unesite naredbu CRL u naredbenu liniju.',
            'Odaberite predloške blokova za kreiranje layouta.',
            'Navedite postavke (naziv layouta, skala, page setup).',
            'Program će kreirati odgovarajuće layoute s postavljenim atributima.',
        ],
        'notes': None,
    },
    {
        'name': 'DBCountV1-1.lsp',
        'category': 'Blokovi i atributi',
        'command': 'DBC',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Napredni brojač blokova koji prebroji sve vrste blokova uključujući dinamičke blokove '
            'u svim stanjima vidljivosti. Rezultate prikazuje u AutoCAD tablici ili izvozi u CSV datoteku.'
        ),
        'usage': [
            'Unesite naredbu DBC u naredbenu liniju.',
            'Odaberite blokove koje želite prebrojati ili koristite sve blokove u nacrtu.',
            'Odaberite format izlaza: tablica u nacrtu ili CSV datoteka.',
            'Za tablicu: odredite točku umetanja u nacrtu.',
            'Za CSV: navedite putanju i naziv datoteke.',
        ],
        'notes': 'Program razlikuje različita stanja vidljivosti dinamičkih blokova kao zasebne vrste.',
    },
    {
        'name': 'ExportLayoutsToExcel.lsp',
        'category': 'Excel integracija',
        'command': 'EXPORTLAYOUTS',
        'author': 'Prilagođeno',
        'description': (
            'Program koji izvozi podatke iz AutoCAD layouta u Excel tablicu. '
            'Prikuplja informacije o layoutima i njihovim atributima te ih organizira '
            'u Excel radnu knjigu za daljnju obradu ili dokumentaciju.'
        ),
        'usage': [
            'Unesite naredbu EXPORTLAYOUTS u naredbenu liniju.',
            'Odaberite layoute za izvoz ili izvezite sve layoute.',
            'Navedite putanju i naziv Excel datoteke.',
            'Program će otvoriti Excel i popuniti podatke.',
        ],
        'notes': 'Potreban je instaliran Microsoft Excel.',
    },
    {
        'name': 'ExtractNestedBlockV1-2.lsp',
        'category': 'Blokovi i atributi',
        'command': 'ENB',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Program koji pretvara ugniježđene (nested) blokove u primarne blokove, zadržavajući '
            'sva svojstva i položaj. Korisno za pojednostavljivanje složenih struktura blokova.'
        ),
        'usage': [
            'Unesite naredbu ENB u naredbenu liniju.',
            'Odaberite blok koji sadrži ugniježđene blokove.',
            'Program će automatski izdvojiti ugniježđene blokove kao primarne.',
        ],
        'notes': None,
    },
    {
        'name': 'GetExcel.lsp',
        'category': 'Excel integracija',
        'command': '(učitava se kao biblioteka)',
        'author': 'R4 by Terry Miller',
        'description': (
            'Biblioteka funkcija za integraciju AutoCAD-a s Microsoft Excelom. Pruža funkcije '
            'za čitanje i pisanje ćelija, otvaranje i zatvaranje radnih knjiga, te rad s listovima. '
            'Koristi se kao osnova za ostale programe koji rade s Excelom (SetFieldsValue, SetLayoutTitles, itd.).'
        ),
        'usage': [
            'Ova datoteka se ne pokreće direktno kao naredba.',
            'Učitava se automatski i pruža funkcije drugim programima.',
            'Programeri mogu koristiti funkcije poput ge:open-excel, ge:read-cell, ge:write-cell.',
        ],
        'notes': 'Potreban je instaliran Microsoft Excel na računalu.',
    },
    {
        'name': 'GrTextV1-1.lsp',
        'category': 'Tekst',
        'command': '(biblioteka funkcija)',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Biblioteka koja vraća grvecs pixel vektorske liste za kodiranje stringova. '
            'Koristi se za prikaz teksta u grafičkom sučelju (npr. preview unutar programa) '
            'i nije namijenjena direktnoj upotrebi od strane krajnjeg korisnika.'
        ),
        'usage': [
            'Ova datoteka se koristi interno od strane ostalih LSP programa.',
            'Ne pokreće se kao samostalna naredba.',
        ],
        'notes': None,
    },
    {
        'name': 'insertBlock.lsp',
        'category': 'Blokovi i atributi',
        'command': 'IB',
        'author': 'Prilagođeno',
        'description': (
            'Program za umetanje blokova i listova s mogućnošću prilagodbe dodjele layera. '
            'Pojednostavljuje postupak umetanja standardnih blokova uz automatsko postavljanje '
            'na odgovarajući layer.'
        ),
        'usage': [
            'Unesite naredbu IB u naredbenu liniju.',
            'Odaberite blok koji želite umetnuti.',
            'Odredite točku umetanja u nacrtu.',
            'Po potrebi odaberite layer za blok.',
        ],
        'notes': None,
    },
    {
        'name': 'InsertElements.lsp',
        'category': 'Blokovi i atributi',
        'command': 'IE',
        'author': 'Prilagođeno',
        'description': (
            'Program za umetanje blokovnih elemenata vezanih uz Maren sustave. '
            'Omogućuje brzo umetanje standardnih komponenti iz definiranih biblioteka blokova.'
        ),
        'usage': [
            'Unesite naredbu IE u naredbenu liniju.',
            'Odaberite element koji želite umetnuti iz ponuđenog izbornika.',
            'Odredite točku umetanja u nacrtu.',
        ],
        'notes': 'Program je prilagođen za specifičnu biblioteku blokova Maren sustava.',
    },
    {
        'name': 'LMpopup.lsp',
        'category': 'Pomoćni alati',
        'command': '(biblioteka funkcija)',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Pomoćna biblioteka koja pruža wrapper za WSH popup metodu za prikaz dijaloških prozora '
            's porukama. Koristi se interno od strane ostalih LSP programa za prikazivanje '
            'obavijesti i upita korisniku.'
        ),
        'usage': [
            'Ova datoteka se koristi interno od strane ostalih LSP programa.',
            'Ne pokreće se kao samostalna naredba.',
        ],
        'notes': None,
    },
    {
        'name': 'makemore.lsp',
        'category': 'Pomoćni alati',
        'command': 'MAKEMORE / MM',
        'author': 'Kent Cooper, Mosad Elewa',
        'description': (
            'Program koji stvara dodatne entitete (kopije) koji odgovaraju svojstvima odabranog objekta. '
            'Podržava oblike, blokove, polilonije, dimenzije i mnoge druge vrste objekata. '
            'Korisno za brzo multipliciranje elemenata s identičnim svojstvima.'
        ),
        'usage': [
            'Unesite naredbu MAKEMORE ili MM u naredbenu liniju.',
            'Odaberite izvorni objekt čija se svojstva žele kopirati.',
            'Odredite točke umetanja za nove kopije.',
            'Pritisnite Enter za završetak.',
        ],
        'notes': None,
    },
    {
        'name': 'MarenConnectPro.lsp',
        'category': 'Solarni sustavi',
        'command': 'MCP',
        'author': 'Prilagođeno',
        'description': (
            'Specijalizirani program za kreiranje dijagrama ožičenja solarnih panela. '
            'Podržava konfiguriranje razmaka i numeriranja za invertere, MPPT-ove i stringove. '
            'Automatski generira sheme spajanja prema zadanim parametrima.'
        ),
        'usage': [
            'Unesite naredbu MCP u naredbenu liniju.',
            'Odaberite konfiguraciju invertera (broj MPPT-ova, stringova po MPPT-u).',
            'Postavite parametre numeriranja i razmake.',
            'Odredite položaj u nacrtu za generiranje dijagrama.',
        ],
        'notes': 'Program je prilagođen za specifične zahtjeve Maren solarnih projekata.',
    },
    {
        'name': 'Mastermind.lsp',
        'category': 'Ostalo',
        'command': 'MASTERMIND',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Implementacija klasične igre Mastermind unutar AutoCAD-a. Igrač ima 8 pokušaja da pogodi '
            'tajni kôd od 4 boje. Nakon svakog pokušaja program prikazuje koliko je boja pogođeno '
            'na ispravnoj poziciji (crni) i koliko na pogrešnoj poziciji (bijeli).'
        ),
        'usage': [
            'Unesite naredbu MASTERMIND u naredbenu liniju.',
            'Program će prikazati sučelje igre u nacrtu.',
            'Odaberite boje za svaki od 4 položaja.',
            'Potvrdite pogađanje i prati povratne informacije.',
            'Pokušajte pogoditi tajni kôd u 8 ili manje pokušaja.',
        ],
        'notes': 'Igra je zamišljena kao demonstracija mogućnosti AutoLisp-a.',
    },
    {
        'name': 'ModifyTextValues.lsp',
        'category': 'Tekst',
        'command': 'MTV',
        'author': 'Prilagođeno',
        'description': (
            'Alat za skupno pretraživanje i zamjenu teksta u nacrtima. Omogućuje zamjenu '
            'tekstualnih vrijednosti s opcijom kontrole smjera obrade (gore-dolje, lijevo-desno). '
            'Radi s Text, MText i atribut objektima.'
        ),
        'usage': [
            'Unesite naredbu MTV u naredbenu liniju.',
            'Unesite tekst za pretraživanje (Find).',
            'Unesite zamjenski tekst (Replace).',
            'Odaberite objekte za pretraživanje ili koristite sve objekte.',
            'Odaberite smjer obrade.',
        ],
        'notes': None,
    },
    {
        'name': 'MTVATT.lsp',
        'category': 'Tekst',
        'command': 'MTVATT',
        'author': 'Prilagođeno',
        'description': (
            'Program za masovnu izmjenu tekstualnih vrijednosti atributa s mogućnošću dodavanja prefiksa '
            'i sufiksa te kontrolom smjera slijeda (povećanje/smanjenje). Koristi se za brzo '
            'preimenovanje ili označavanje serija blokova u elektroenergetskim i industrijskim shemama.'
        ),
        'usage': [
            'Unesite naredbu MTVATT u naredbenu liniju.',
            'Odaberite blokove s atributima koje želite izmijeniti.',
            'Unesite prefiks i/ili sufiks koji se dodaju vrijednostima.',
            'Odaberite smjer slijeda (inkrementalni ili dekrementalni).',
            'Potvrdite izmjene.',
        ],
        'notes': None,
    },
    {
        'name': 'NumIncV4-1.lsp',
        'category': 'Tekst',
        'command': 'NI / NUMINC',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Sveobuhvatan alat za inkrementalno numeriranje koji omogućuje postavljanje uzastopnih '
            'oznaka (brojeva ili slova) u nacrt. Podržava prilagodbu formata, korak povećanja, '
            'predloške i array mogućnosti za brzo stvaranje nizova oznaka.'
        ),
        'usage': [
            'Unesite naredbu NI ili NUMINC u naredbenu liniju.',
            'Postavite parametre: početna vrijednost, korak, format, predložak.',
            'Kliknite na mjesta u nacrtu gdje želite postaviti numeriranu oznaku.',
            'Program automatski povećava vrijednost za svaki klik.',
            'Pritisnite Enter ili Esc za završetak.',
        ],
        'notes': 'Podržava kombiniranje teksta i brojeva u oznakama (npr. "PAN-001", "PAN-002"...).',
    },
    {
        'name': 'panelRenum.lsp',
        'category': 'Solarni sustavi',
        'command': 'PANELRENUM',
        'author': 'Prilagođeno',
        'description': (
            'Program za preimenovanje blokova solarnih panela prema shemi numeriranja '
            'inverter/MPPT/string/panel s live pregledom u nacrtu. '
            'Automatizira oznakavanje panela u solarnim projektima prema hijerarhijskoj strukturi sustava.'
        ),
        'usage': [
            'Unesite naredbu PANELRENUM u naredbenu liniju.',
            'Odaberite blokove solarnih panela za preimenovanje.',
            'Postavite parametre numeriranja (inverter, MPPT, string, panel).',
            'Pregledajte live pregled u nacrtu.',
            'Potvrdite primjenu numeriranja.',
        ],
        'notes': None,
    },
    {
        'name': 'RenumberLayoutsV1-2.lsp',
        'category': 'Layouti',
        'command': 'RNL',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Program koji automatski renumerira sve paperspace layoute u nacrtu s opcionalnim '
            'prefiksom i sufiksom. Korisno za ujednačavanje redosljeda i oznaka layouta '
            'u projektima s više listova.'
        ),
        'usage': [
            'Unesite naredbu RNL u naredbenu liniju.',
            'Po želji unesite prefiks (npr. "List-") i/ili sufiks (npr. "-A").',
            'Unesite početni broj.',
            'Program će renumerirati sve layoute prema zadanom redoslijedu.',
        ],
        'notes': 'Redoslijed numeriranja odgovara redoslijedu kartica layouta.',
    },
    {
        'name': 'setCBelements.lsp',
        'category': 'Električni nacrti',
        'command': 'SETCBE',
        'author': 'Prilagođeno',
        'description': (
            'Interaktivni alat za postavljanje opisa elemenata sklopnih blokova s unaprijed definiranim '
            'listama odabira (snage, karakteristike, tipovi razvodnih uređaja, faze). '
            'Olakšava standardizirano opisivanje zaštitnih komponenti u elektroenergetskim shemama.'
        ),
        'usage': [
            'Unesite naredbu SETCBE u naredbenu liniju.',
            'Odaberite blok sklopnog elementa u nacrtu.',
            'Iz padajućih lista odaberite: snagu, karakteristiku, tip i broj faza.',
            'Potvrdite odabir – program će ažurirati atribute bloka.',
        ],
        'notes': 'Liste odabira su prilagođene za standardne komponente prema projektnim specifikacijama.',
    },
    {
        'name': 'setLayoutOrderV5.lsp',
        'category': 'Layouti',
        'command': 'SLO',
        'author': 'Prilagođeno',
        'description': (
            'Program koji postavlja redosljed layouta korištenjem interaktivnog sustava numeriranja. '
            'Omogućuje vizualno upravljanje redoslijedom kartica layouta u nacrtu.'
        ),
        'usage': [
            'Unesite naredbu SLO u naredbenu liniju.',
            'Prikazat će se lista svih layouta.',
            'Dodijelite numeričke vrijednosti za željeni redoslijed.',
            'Potvrdite – program će preurediti kartice layouta.',
        ],
        'notes': None,
    },
    {
        'name': 'SetFieldsValue.lsp',
        'category': 'Excel integracija',
        'command': 'SFV',
        'author': 'Prilagođeno',
        'description': (
            'Program koji ažurira svojstva nacrta (drawing properties) iz podataka u Excel tablici. '
            'Automatizira popunjavanje naslova, broja projekta i ostalih informacija titlebloka '
            'iz centralne Excel baze podataka.'
        ),
        'usage': [
            'Unesite naredbu SFV u naredbenu liniju.',
            'Odaberite Excel datoteku s podacima.',
            'Program će pročitati odgovarajuće ćelije i ažurirati drawing properties nacrta.',
            'Ažurirani fields se automatski osvježavaju u titlebloku.',
        ],
        'notes': 'Excel datoteka mora imati propisanu strukturu s definiranim ćelijama za svako polje.',
    },
    {
        'name': 'SetLayoutTitles.lsp',
        'category': 'Layouti',
        'command': 'SLT',
        'author': 'Prilagođeno',
        'description': (
            'Program koji postavlja naslove layouta i vrijednosti skale iz Excel tablice. '
            'Ubrzava postavljanje titleblokovskih informacija za projekte s velikim brojem listova.'
        ),
        'usage': [
            'Unesite naredbu SLT u naredbenu liniju.',
            'Odaberite Excel datoteku s popisom layouta, naslovima i skalama.',
            'Program će automatski ažurirati odgovarajuće atribute titlebloka na svakom layoutu.',
        ],
        'notes': 'Excel datoteka mora sadržavati stupce: naziv layouta, naslov, skala.',
    },
    {
        'name': 'simpleCount.lsp',
        'category': 'Blokovi i atributi',
        'command': 'SC',
        'author': 'Prilagođeno',
        'description': (
            'Jednostavan alat za prebrojavanje instanci odabranog bloka u nacrtu. '
            'Prikazuje rezultate u dijalogu ili direktno u naredbenom retku. '
            'Brz i praktičan za brze provjere količine elemenata.'
        ),
        'usage': [
            'Unesite naredbu SC u naredbenu liniju.',
            'Odaberite blokove koje želite prebrojati.',
            'Program će prikazati broj pronađenih instanci.',
        ],
        'notes': None,
    },
    {
        'name': 'StealV1-8.lsp',
        'category': 'Blokovi i atributi',
        'command': 'STEAL',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Program koji uvozi blokove, layere, linetypes, stilove i ostale komponente nacrta '
            'iz drugog DWG dokumenta. Korisno za prenošenje standardiziranih elemenata između projekata '
            'bez ručnog kopiranja ili otvaranja izvorne datoteke.'
        ),
        'usage': [
            'Unesite naredbu STEAL u naredbenu liniju.',
            'Odaberite izvornu DWG datoteku iz koje želite uvesti elemente.',
            'Odaberite kategoriju elemenata: blokovi, layeri, linetypes, stilovi teksta itd.',
            'Odaberite konkretne elemente za uvoz.',
            'Program će uvesti odabrane elemente u trenutni nacrt.',
        ],
        'notes': 'Ako element s istim imenom već postoji, program će upitati o prepisivanju.',
    },
    {
        'name': 'sumPowerV2.lsp',
        'category': 'Solarni sustavi',
        'command': 'SUMPOWER',
        'author': 'Prilagođeno',
        'description': (
            'Program koji izračunava ukupnu snagu iz odabranih tekstualnih objekata, izračunava '
            'korigiranu snagu s faktorom sličnosti (similarity factor) te struju. '
            'Koristi se za brze proračune u solarnim i elektroenergetskim projektima.'
        ),
        'usage': [
            'Unesite naredbu SUMPOWER u naredbenu liniju.',
            'Odaberite tekstualne objekte koji sadrže vrijednosti snage (u kW ili W).',
            'Unesite faktor sličnosti (0–1).',
            'Program će prikazati: ukupnu snagu, korigiranu snagu i izračunatu struju.',
        ],
        'notes': 'Vrijednosti snage u tekstu moraju biti u numeričkom formatu s jedinicom (npr. "5.2 kW").',
    },
    {
        'name': 'TabSortV2-2.lsp',
        'category': 'Layouti',
        'command': 'TABSORT',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Napredni upravljač karticama layouta s mogućnostima sortiranja, preimenovanja, kopiranja '
            'i funkcijom traženja i zamjene. Pruža grafičko sučelje za jednostavno upravljanje '
            'redoslijedom i imenima layouta u kompleksnim projektima.'
        ),
        'usage': [
            'Unesite naredbu TABSORT u naredbenu liniju.',
            'Otvorit će se dijaloški prozor s listom svih layouta.',
            'Koristite drag & drop ili gumbe za promjenu redoslijeda.',
            'Koristite "Rename" za preimenovanje layouta.',
            'Koristite "Find & Replace" za skupno preimenovanje.',
            'Pritisnite "OK" za primjenu promjena.',
        ],
        'notes': None,
    },
    {
        'name': 'TCountV1-1.lsp',
        'category': 'Tekst',
        'command': 'TC',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Program koji prebroji pojave određenog stringa u odabranim objektima i prikazuje '
            'rezultate u obliku AutoCAD tablice. Podržava pretraživanje po tekst, MText i atribut objektima.'
        ),
        'usage': [
            'Unesite naredbu TC u naredbenu liniju.',
            'Unesite string koji se traži.',
            'Odaberite objekte za pretraživanje.',
            'Odredite točku umetanja tablice s rezultatima.',
        ],
        'notes': None,
    },
    {
        'name': 'Text2MTextV2-0.lsp',
        'category': 'Tekst',
        'command': 'T2MT',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Program koji pretvara jednolinjski tekst (Text) u višelinjski tekst (MText) s fleksibilnim '
            'opcijama postavljanja. Zadržava sva vizualna svojstva originalnog teksta i nudi '
            'mogućnost grupiranja više text objekata u jedan MText.'
        ),
        'usage': [
            'Unesite naredbu T2MT u naredbenu liniju.',
            'Odaberite Text objekte koje želite pretvoriti.',
            'Odaberite opciju postavljanja (na istom mjestu, grupirano itd.).',
            'Program će zamijeniti odabrane Text objekte s MText ekvivalentima.',
        ],
        'notes': 'Original Text objekti se brišu po završetku konverzije.',
    },
    {
        'name': 'tlen.lsp',
        'category': 'Pomoćni alati',
        'command': 'TLEN',
        'author': 'Lee Mac (www.lee-mac.com)',
        'description': (
            'Program koji prikazuje ukupnu duljinu odabranih geometrijskih objekata '
            '(lukovi, kružnice, linije, polilonije, splajnovi). '
            'Korisno za brzo mjerenje ukupne dužine trase ili skupa elemenata.'
        ),
        'usage': [
            'Unesite naredbu TLEN u naredbenu liniju.',
            'Odaberite geometrijske objekte (linije, polilonije, lukove itd.).',
            'Pritisnite Enter za potvrdu selekcije.',
            'Program će prikazati ukupnu duljinu odabranih objekata.',
        ],
        'notes': 'Rezultat se prikazuje u trenutnim jedinicama nacrta.',
    },
    {
        'name': 'VDS.lsp',
        'category': 'Električni nacrti',
        'command': 'VDS',
        'author': 'Prilagođeno',
        'description': (
            'Program za ažuriranje AutoWire atributa u VDS (Visual Distribution System) shemama. '
            'Automatski slijedi lančanu strukturu i numerira elemente. '
            'Koristi se za automatizaciju označavanja razvoda u elektroenergetskim projektima.'
        ),
        'usage': [
            'Unesite naredbu VDS u naredbenu liniju.',
            'Odaberite početni element VDS lanca.',
            'Program će automatski pratiti lančanu strukturu i ažurirati atribute.',
            'Potvrdite numeriranje.',
        ],
        'notes': 'Program pretpostavlja propisanu strukturu blokova prema AutoWire standardu.',
    },
    {
        'name': 'VDSnum2.lsp',
        'category': 'Električni nacrti',
        'command': 'VDSNUMPRO / VDSCHECKPRO',
        'author': 'Prilagođeno',
        'description': (
            'Napredni alat za VDS numeriranje (VDSNUMPRO) s live pregledom u nacrtu te alatom '
            'za provjeru dupliciranih oznaka i praznina u numeriranju (VDSCHECKPRO). '
            'Pruža potpunu kontrolu nad numeriranjem VDS elemenata u projektima razvoda.'
        ),
        'usage': [
            'Za numeriranje: unesite naredbu VDSNUMPRO.',
            'Postavite parametre numeriranja i pratite live pregled.',
            'Za provjeru: unesite naredbu VDSCHECKPRO.',
            'Program će prikazati duplikate i praznine u numeriranju s vizualnim markacama.',
        ],
        'notes': 'VDSCHECKPRO je posebno koristan za provjeru ispravnosti numeriranja na kraju projekta.',
    },
    {
        'name': 'VportsToggle.lsp',
        'category': 'Layouti',
        'command': 'VT',
        'author': 'Prilagođeno',
        'description': (
            'Program koji prebacuje između jednog i dva vertikalna viewporta u model spaceu. '
            'Korisno za uspoređivanje različitih dijelova nacrta istovremeno.'
        ),
        'usage': [
            'Unesite naredbu VT u naredbenu liniju.',
            'Program će prebaciti izgled: s jednog viewporta na dva vertikalna ili obratno.',
        ],
        'notes': None,
    },
]

# Sortiraj abecedno
programs.sort(key=lambda x: x['name'].lower())

for prog in programs:
    doc.add_heading(prog['name'], level=2)
    add_horizontal_line(doc)
    doc.add_paragraph()

    add_info_row(doc, 'Kategorija', prog['category'])
    add_info_row(doc, 'Naredba', prog['command'])
    add_info_row(doc, 'Autor', prog['author'])
    doc.add_paragraph()

    desc_heading = doc.add_paragraph()
    desc_heading_run = desc_heading.add_run('Opis')
    desc_heading_run.bold = True
    desc_heading_run.font.size = Pt(11)
    desc_heading_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    desc_p = doc.add_paragraph(prog['description'])
    desc_p.paragraph_format.space_after = Pt(6)

    usage_heading = doc.add_paragraph()
    usage_heading_run = usage_heading.add_run('Upute za korištenje')
    usage_heading_run.bold = True
    usage_heading_run.font.size = Pt(11)
    usage_heading_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    for step_i, step in enumerate(prog['usage'], 1):
        step_p = doc.add_paragraph(style='List Number')
        step_p.text = ''
        step_p.paragraph_format.space_before = Pt(1)
        step_p.paragraph_format.space_after = Pt(1)
        run = step_p.add_run(step)
        run.font.size = Pt(10.5)

    if prog.get('notes'):
        doc.add_paragraph()
        note_p = doc.add_paragraph()
        note_bold = note_p.add_run('Napomena: ')
        note_bold.bold = True
        note_bold.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        note_text = note_p.add_run(prog['notes'])
        note_text.font.size = Pt(10.5)
        note_p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()

# ==================== POPIS NAREDBI ====================
doc.add_page_break()
doc.add_heading('Brzi pregled naredbi', level=1)

intro_cmd = doc.add_paragraph(
    'U tablici ispod navedene su sve raspoložive naredbe za brzu referencu.'
)
intro_cmd.paragraph_format.space_after = Pt(10)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
headers = ['Program', 'Naredba', 'Kategorija']
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    run = hdr_cells[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F497D')
    hdr_cells[i]._tc.get_or_add_tcPr().append(shd)

fill_colors = ['F2F7FC', 'FFFFFF']
for idx, prog in enumerate(programs):
    row_cells = table.add_row().cells
    row_cells[0].text = prog['name']
    row_cells[1].text = prog['command']
    row_cells[2].text = prog['category']
    fill = fill_colors[idx % 2]
    for cell in row_cells:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        cell._tc.get_or_add_tcPr().append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

for col_idx, width in enumerate([Cm(6), Cm(5), Cm(5)]):
    for row in table.rows:
        row.cells[col_idx].width = width

output_path = '/home/user/AutoLisp/AutoLisp_Dokumentacija.docx'
doc.save(output_path)
print(f'Dokument spremljen: {output_path}')
